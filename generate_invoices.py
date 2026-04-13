"""
generate_invoices_py.py
Pure-Python replacement for generate_invoices.js
Requires: python-docx  (pip install python-docx)
"""

import json, os, math
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy

# ── Constants ────────────────────────────────────────────────────────────────
MONTH       = "Feb- 2026"
DATE        = "28-02-2026"
MONTH_LABEL = "Feb'26"
MONTHLY     = 16500
MONTH_DAYS  = 28
SIG_DIR     = "signatures/"

SIG_MAP = {
  "Ahmed Azam":                   "Ahmed_Azam-removebg-preview.png",
  "Anand Ganesh  Chopade":        "Anand Ganesh Chopade.png",
  "Arundathi Jalagam":            "Arundathi_Jalagam-removebg-preview.png",
  "Bonagiri Rehana":              "Bonagiri_Rehana-removebg-preview.png",
  "Chouta Keerthana":             "Chouta_Keerthana-removebg-preview.png",
  "Deshapaga Raghavendar":        "Deshapaga_Raghavendar-removebg-preview.png",
  "Dhenumakonda Lavanya":         "Dhenumakonda_Lavanya-removebg-preview.png",
  "Diravath Mounika":             "Diravath_Mounika-removebg-preview.png",
  "Gaja Bala Narayana":           "Gaja Bala Narayana-removebg-preview.png",
  "Gopala Saritha":               "Gopala_Saritha-removebg-preview.png",
  "Gudise Hemanth Kumar":         "Gudise_Hemanth_Kumar-removebg-preview.png",
  "Jetti Hima Sindhu":            "Jetti_Hima_Sindhu-removebg-preview.png",
  "K THIRUPATHAIAH":              "K_THIRUPATHAIAH-removebg-preview.png",
  "Kanche Srisailam":             "Kanche_Srisailam-removebg-preview.png",
  "Kasireddy Harish":             "Kasireddy_Harish-removebg-preview.png",
  "Kasturi Sathish":              "Kasturi_Sathish-removebg-preview.png",
  "KATRAVATH MANGESH":            "KATRAVATH_MANGESH-removebg-preview.png",
  "Katravath Mohan Rathod":       "Katravath_Mohan_Rathod-removebg-preview.png",
  "Katravath Radhika":            "Katravath_Radhika-removebg-preview.png",
  "Kodavath Swamy":               "Kodavath_Swamy-removebg-preview.png",
  "KOLLA SUDHAKAR":               "KOLLA_SUDHAKAR-removebg-preview.png",
  "Kunduru Upender Reddy":        "Kunduru_Upender_Reddy-removebg-preview.png",
  "M JATHIN SAI":                 "M_JATHIN_SAI-removebg-preview.png",
  "M Sunitha":                    "M_SUNITHA-removebg-preview.png",
  "M Swarupa":                    "M_Swarupa-removebg-preview.png",
  "Meka Esthar Rani":             "Meka Esthar Rani-removebg-preview.png",
  "Mamidi Raj kumar":             "Mamidi_Raj_kumar-removebg-preview.png",
  "Md Mainoddin":                 "Md_Mainoddin-removebg-preview.png",
  "Mekala Anusha":                "MEKALA_ANUSHA-removebg-preview.png",
  "MOHAMMAD NAZIYA":              "MOHAMMAD_NAZIYA-removebg-preview.png",
  "Mohammad Shaheen Begum":       "Mohammad Shaheen Begum-removebg-preview.png",
  "Mohammed Yaqoob khan":         "Mohammed_Yaqoob_Khan-removebg-preview.png",
  "Mudavath Balakoti":            "Mudavath_Balakoti-removebg-preview.png",
  "MUDAVATH KIRAN":               "MUDAVATH_KIRAN-removebg-preview.png",
  "Mudavath Ramesh":              "Mudavath_Ramesh-removebg-preview.png",
  "Neeli Sreevani":               "Neeli SreevaniSignature remove.png",
  "Padira Radhika":               "Padira_Radhika-removebg-preview.png",
  "Pandiri Punith kumar":         "Pandiri_Punith_kumar-removebg-preview.png",
  "Poojari Ramu":                 "Pujari_Ramu-removebg-preview.png",
  "Porandla Chandar":             "Porandla_Chandar-removebg-preview.png",
  "Ramavath Saimahesh Nayak":     "Ramavath_Saimahesh_Nayak-removebg-preview.png",
  "Ramavath Uma Mahesh":          "Ramavath_Uma_Mahesh-removebg-preview.png",
  "Ranabotu Saidi Reddy":         "Ranabotu_Saidi_Reddy-removebg-preview.png",
  "Sadde Sindhura":               "Sadde_Sindhura-removebg-preview.png",
  "Shaikh abdul Avesh":           "Shaikh_abdul_Avesh-removebg-preview.png",
  "Shivarathri Swapna":           "Shivarathri_Swapna-removebg-preview.png",
  "Tandra Sabastin":              "Tandra_Sabastin-removebg-preview.png",
  "Tabassum Afreen":              "Thabasum_Afreen-removebg-preview.png",
  "Thatipally Manoj Kumar":       "Thatipally_Manoj_Kumar-removebg-preview.png",
  "Thurpati vijay Baskar":        "Thurpati_vijay_Baskar-removebg-preview.png",
  "Ushanolla Ravali":             "Ushanolla_Ravali-removebg-preview.png",
  "Ushanula Ramya":               "Ushanula_Ramya-removebg-preview.png",
  "Chejerla Nagavamsidhar Reddy": "Chejerla_Nagavamsidhar_Reddy-removebg-preview.png",
}

# ── Helpers ───────────────────────────────────────────────────────────────────
def rupees(n):
    n = round(n)
    s = str(n)
    if len(s) <= 3:
        return s
    result = s[-3:]
    s = s[:-3]
    while len(s) > 2:
        result = s[-2:] + ',' + result
        s = s[:-2]
    if s:
        result = s + ',' + result
    return result


def get_sig_path(emp):
    name_key = emp.get('service_provider') or emp.get('name', '')
    filename = SIG_MAP.get(name_key)
    if not filename:
        for k, v in SIG_MAP.items():
            if k.strip().lower() == name_key.strip().lower():
                filename = v
                break
    if not filename:
        filename = emp.get('sig_filename')
    if not filename:
        return None
    for candidate in [os.path.join(SIG_DIR, filename), filename]:
        if os.path.exists(candidate):
            return candidate
    print(f"  [WARN] Signature not found for \"{name_key}\" (tried: {filename})")
    return None


def set_cell_border(cell, color="000000"):
    """Add a thin black border to a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def set_cell_shading(cell, fill_hex):
    """Set background shading of a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)


def cell_para(cell, text, bold=False, sz=10, align=WD_ALIGN_PARAGRAPH.LEFT):
    """Write a paragraph into a table cell, clearing existing content."""
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(sz)
    run.font.name = 'Arial'
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_para(doc, text_parts, space_after=Pt(4)):
    """
    text_parts: list of (text, bold) tuples
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = Pt(0)
    for text, bold in text_parts:
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(10)
        run.font.name = 'Arial'
    return p


def set_col_width(table, col_idx, width_cm):
    """Force a column width."""
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)


# ── Page builder ───────────────────────────────────────────────────────────────
def build_page(doc, emp, add_page_break=False):
    rate = MONTHLY / MONTH_DAYS
    total_amt = emp['total_amount']
    projs = emp.get('projects', [])

    sp = Pt(3)  # small spacing after each paragraph

    add_para(doc, [(f"Date: - {DATE}", True)], sp)
    add_para(doc, [("", False)], sp)
    add_para(doc, [("TO,", True)], sp)
    add_para(doc, [("CUBE Highways Technologies Private Limited,", True)], sp)
    add_para(doc, [("3rd Floor, GMR Aero Towers – 2,", False)], sp)
    add_para(doc, [("Mamidipally Village, Saroor Nagar Mandal,", False)], sp)
    add_para(doc, [("Ranga Reddy, Hyderabad, Telangana - 500108", False)], sp)
    add_para(doc, [("", False)], sp)
    add_para(doc, [("GST No- ", True), ("36AAKCC7533R1ZW", False)], sp)
    add_para(doc, [("PAN No- ", True), ("AAKCC7533R", False)], sp)
    add_para(doc, [("", False)], sp)
    add_para(doc, [("Sir,", True)], sp)
    add_para(doc, [("", False)], sp)
    add_para(doc, [
        ("Subject: ", True),
        ("Consultant fee for ", False),
        (MONTH, True),
        (" data processing & Analysis ", False),
        (f"Rs.{rupees(total_amt)}/-", True),
        (" per month. The commercials are mentioned below.", False),
    ], Pt(6))

    # ── Fee table ──────────────────────────────────────────────────────────────
    # Col widths in cm: Particulars | Days | WBS | Amount
    col_w = [6.5, 3.2, 6.5, 3.5]
    fee_table = doc.add_table(rows=0, cols=4)
    fee_table.style = 'Table Grid'

    # Header row
    hdr = fee_table.add_row()
    headers = [
        f"Particulars",
        f"No. of working\ndays in {MONTH_LABEL}",
        "WBS Elements",
        "Payable Amount\n(Rs.)"
    ]
    for i, (cell, txt) in enumerate(zip(hdr.cells, headers)):
        cell_para(cell, txt, bold=True, sz=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(cell, "BDD7EE")
        set_cell_border(cell)

    # Data rows
    if not projs:
        row = fee_table.add_row()
        cell_para(row.cells[0], "Consultant fee – Data Processing", sz=9)
        cell_para(row.cells[1], str(emp['attendance']), sz=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        cell_para(row.cells[2], "", sz=9)
        cell_para(row.cells[3], rupees(total_amt), sz=9, align=WD_ALIGN_PARAGRAPH.RIGHT)
        for c in row.cells: set_cell_border(c)
        for _ in range(2):
            r = fee_table.add_row()
            for c in r.cells:
                cell_para(c, "", sz=9)
                set_cell_border(c)
    else:
        for idx, proj in enumerate(projs):
            amt = round(proj['days'] * rate)
            row = fee_table.add_row()
            cell_para(row.cells[0], "Consultant fee – Data Processing" if idx == 0 else "", sz=9)
            cell_para(row.cells[1], str(proj['days']), sz=9, align=WD_ALIGN_PARAGRAPH.CENTER)
            cell_para(row.cells[2], proj.get('wbs', ''), sz=9)
            cell_para(row.cells[3], rupees(amt), sz=9, align=WD_ALIGN_PARAGRAPH.RIGHT)
            for c in row.cells: set_cell_border(c)
        # Pad to at least 3 data rows
        while len(fee_table.rows) < 4:
            r = fee_table.add_row()
            for c in r.cells:
                cell_para(c, "", sz=9)
                set_cell_border(c)

    # Total row
    tr = fee_table.add_row()
    cell_para(tr.cells[0], "", sz=9)
    cell_para(tr.cells[1], str(emp['attendance']), bold=True, sz=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_para(tr.cells[2], "Total Pay", bold=True, sz=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_para(tr.cells[3], f"Rs. {rupees(total_amt)}/-", bold=True, sz=9, align=WD_ALIGN_PARAGRAPH.RIGHT)
    set_cell_shading(tr.cells[2], "E2EFDA")
    set_cell_shading(tr.cells[3], "E2EFDA")
    for c in tr.cells: set_cell_border(c)

    # Apply column widths
    for row in fee_table.rows:
        for i, w in enumerate(col_w):
            row.cells[i].width = Cm(w)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    add_para(doc, [("Thanking you and always assuring you of our best services.", False)], sp)
    add_para(doc, [("", False)], sp)
    add_para(doc, [("Yours faithfully", True)], sp)
    add_para(doc, [("", False)], sp)
    add_para(doc, [("Authorised Signature", True)], sp)

    # Signature image
    sig_path = get_sig_path(emp)
    sig_para = doc.add_paragraph()
    sig_para.paragraph_format.space_after = Pt(4)
    if sig_path:
        try:
            run = sig_para.add_run()
            run.add_picture(sig_path, width=Cm(4.5), height=Cm(1.8))
        except Exception as e:
            print(f"  [WARN] Could not embed signature for {emp.get('service_provider','')}: {e}")

    add_para(doc, [("Service Provider: ", True), (emp.get('service_provider') or emp.get('name', ''), False)], sp)
    add_para(doc, [("Address: ", True), (emp.get('address', ''), False)], sp)
    add_para(doc, [("Email- ", True), (emp.get('email', ''), False)], sp)
    add_para(doc, [("Contact No. ", True), (emp.get('contact', ''), False)], sp)
    add_para(doc, [("PAN No- ", True), (emp.get('pan', ''), False)], sp)
    add_para(doc, [("", False)], sp)
    add_para(doc, [("Bank details below:", True)], sp)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ── Bank table ─────────────────────────────────────────────────────────────
    bank_table = doc.add_table(rows=2, cols=4)
    bank_table.style = 'Table Grid'
    bank_hdrs = ["Account-Name", "Bank Name", "Bank Account Number", "IFSC Code"]
    bank_vals = [
        emp.get('account_name', ''),
        emp.get('bank_name', ''),
        str(emp.get('account_number', '')),
        emp.get('ifsc', ''),
    ]
    bank_col_w = [4.9, 4.9, 5.4, 4.5]
    for i, txt in enumerate(bank_hdrs):
        c = bank_table.rows[0].cells[i]
        cell_para(c, txt, bold=True, sz=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(c, "BDD7EE")
        set_cell_border(c)
    for i, txt in enumerate(bank_vals):
        c = bank_table.rows[1].cells[i]
        cell_para(c, txt, sz=9)
        set_cell_border(c)
    for row in bank_table.rows:
        for i, w in enumerate(bank_col_w):
            row.cells[i].width = Cm(w)

    if add_page_break:
        doc.add_page_break()


# ── Main ───────────────────────────────────────────────────────────────────────
def main():
    employees = json.load(open('emp_data.json', encoding='utf-8'))

    doc = Document()

    # Page setup: Letter-ish wide margins
    section = doc.sections[0]
    section.page_width  = Cm(21.59)
    section.page_height = Cm(27.94)
    section.top_margin    = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin   = Cm(1.8)
    section.right_margin  = Cm(1.8)

    # Default paragraph font
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)

    for idx, emp in enumerate(employees):
        build_page(doc, emp, add_page_break=(idx < len(employees) - 1))
        print(f"  Built page for {emp.get('service_provider') or emp.get('name','')}")

    out = 'Employee_Invoices_new.docx'
    doc.save(out)
    print(f"\nDone: {out}  ({len(employees)} invoices)")


if __name__ == '__main__':
    main()